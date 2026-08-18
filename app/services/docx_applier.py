"""docx 导出服务（require.md 3.3：基于原文档生成修订版）。

仅应用「已接受 / 已自行修改」的建议：
- 文本类（typo/term/polish）：同段内多条建议按「原文长度降序 + 区间占用跟踪」合并替换，
  避免重叠建议相互破坏（如整句优化与其中错别字修复并存）
- 格式类（format）：按 doc_position 中结构化信息（format_field/expected）修改段落格式

注：v0.1 段落级替换会合并该段 run 的字体设置，复杂混合格式段落建议人工复核。
"""
from __future__ import annotations

import logging
import shutil
from collections import defaultdict

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.common import SuggestionStatus

logger = logging.getLogger(__name__)

_ALIGN_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_suggestions_to_docx(
    source_path: str,
    suggestions: list,
    output_path: str,
) -> str:
    """把建议应用到源 docx 的副本上，输出到 output_path。"""
    shutil.copyfile(source_path, output_path)
    doc = DocxDocument(output_path)

    paragraphs = _all_paragraphs(doc)

    # 按段落分组
    text_by_para: dict[int, list] = defaultdict(list)
    format_by_para: dict[int, list] = defaultdict(list)

    for sug in suggestions:
        if sug.status not in (SuggestionStatus.ACCEPTED.value, SuggestionStatus.MODIFIED.value):
            continue
        pos = sug.doc_position or {}
        para_idx = pos.get("paragraph_index")
        if para_idx is None or para_idx >= len(paragraphs):
            continue
        if sug.review_type == "format":
            format_by_para[para_idx].append(sug)
        else:
            text_by_para[para_idx].append(sug)

    applied = 0
    for para_idx in sorted(set(text_by_para) | set(format_by_para)):
        paragraph = paragraphs[para_idx]
        applied += _apply_text_suggestions(paragraph, text_by_para.get(para_idx, []))
        for sug in format_by_para.get(para_idx, []):
            if _apply_format_fix(paragraph, sug.doc_position or {}):
                applied += 1

    doc.save(output_path)
    logger.info("导出 %s：应用 %d/%d 条建议", output_path, applied, len(suggestions))
    return output_path


def _all_paragraphs(doc) -> list:
    """文档全部段落（含表格内单元格段落），与解析器顺序一致。"""
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras


def _apply_text_suggestions(paragraph, suggestions: list) -> int:
    """同段文本建议合并替换：原文长度降序优先，区间不重叠。

    对每个建议，在原文本中寻找「未被其他建议占用」的第一个出现位置；
    全部定位完成后基于原文本一次性重建（避免顺序替换破坏后续原文匹配）。
    """
    if not suggestions:
        return 0
    full = paragraph.text

    items = sorted(
        (
            (s.original_text, s.modified_text or s.suggested_text)
            for s in suggestions
            if s.original_text
        ),
        key=lambda x: len(x[0]),
        reverse=True,  # 长原文优先，覆盖嵌套片段
    )

    occupied: list[tuple[int, int]] = []
    replacements: list[tuple[int, int, str]] = []

    for original, new in items:
        found = None
        start = 0
        while True:
            pos = full.find(original, start)
            if pos == -1:
                break
            end = pos + len(original)
            if not any(not (end <= s or pos >= e) for s, e in occupied):
                found = (pos, end)
                break
            start = pos + 1
        if found is None:
            continue
        occupied.append(found)
        replacements.append((found[0], found[1], new))

    if not replacements:
        return 0

    # 基于原文本一次性重建
    parts: list[str] = []
    last = 0
    for pos, end, new in sorted(replacements):
        parts.append(full[last:pos])
        parts.append(new)
        last = end
    parts.append(full[last:])
    new_text = "".join(parts)

    font_info = None
    if paragraph.runs:
        r0 = paragraph.runs[0]
        font_info = (r0.font.name, r0.font.size, r0.font.bold)
    for run in paragraph.runs:
        run.text = ""
    r = paragraph.add_run(new_text)
    if font_info:
        r.font.name, r.font.size, r.font.bold = font_info
    return len(replacements)


def _apply_format_fix(paragraph, pos: dict) -> bool:
    field = pos.get("format_field")
    expected = str(pos.get("expected", ""))
    if not field or not expected:
        return False
    pf = paragraph.paragraph_format

    if field == "字体":
        for run in paragraph.runs:
            run.font.name = expected
        return True
    if field == "字号":
        size_pt = _parse_float(expected)
        if size_pt is None:
            return False
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
        return True
    if field == "首行缩进":
        chars = _parse_float(expected)
        if chars is None:
            return False
        size = None
        for run in paragraph.runs:
            if run.font.size:
                size = run.font.size.pt
                break
        pf.first_line_indent = Pt(chars * (size or 12.0))
        return True
    if field == "行距":
        if "磅" in expected:
            pt = _parse_float(expected)
            if pt is None:
                return False
            pf.line_spacing = Pt(pt)
        else:
            mult = _parse_float(expected)
            if mult is None:
                return False
            pf.line_spacing = mult
        return True
    if field == "对齐方式":
        align = _ALIGN_MAP.get(expected)
        if align is None:
            return False
        pf.alignment = align
        return True
    if field == "加粗":
        for run in paragraph.runs:
            run.font.bold = True
        return True
    return False


def _parse_float(s: str) -> float | None:
    try:
        return float("".join(c for c in s if c.isdigit() or c == "."))
    except (ValueError, TypeError):
        return None
