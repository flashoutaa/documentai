"""docx 解析服务（require.md 5.4 文档处理管线）。

解析 .docx 为结构化段落：
- 文本内容
- 样式信息：字体、字号、首行缩进、行距、对齐、加粗
- 段落类型推断（标题/正文/落款）
"""
from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.chunker import chunk_paragraphs


@dataclass
class ParagraphMeta:
    """单段落的结构化信息。"""

    index: int  # 文档内段落序号（0 起）
    text: str
    font: str | None = None
    size_pt: float | None = None
    first_line_indent_chars: float | None = None  # 首行缩进（字符）
    line_spacing: float | None = None  # 倍数
    line_spacing_pt: float | None = None  # 固定行距（磅）
    space_before_pt: float | None = None
    space_after_pt: float | None = None
    alignment: str | None = None  # left|center|right|justify
    bold: bool | None = None
    style_name: str | None = None  # Word 内置样式名（Heading 1 等）
    is_heading: bool = False
    heading_level: int | None = None  # 1-9


@dataclass
class ParsedDoc:
    paragraphs: list[ParagraphMeta] = field(default_factory=list)
    char_count: int = 0

    def body_texts(self) -> list[str]:
        return [p.text for p in self.paragraphs]


def _alignment_str(alignment) -> str | None:
    if alignment is None:
        return None
    return {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }.get(alignment)


def _run_font(p) -> tuple[str | None, float | None, bool | None]:
    """取段落首个 run 的字体/字号/加粗（正文级格式通常整段一致）。"""
    font_name = size_pt = bold = None
    for run in p.runs:
        if run.text.strip():
            if run.font.name:
                font_name = run.font.name
            if run.font.size is not None:
                size_pt = run.font.size.pt
            bold = run.bold
            break
    return font_name, size_pt, bold


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    lowered = style_name.lower()
    if "heading" in lowered:
        digits = "".join(c for c in lowered if c.isdigit())
        return int(digits) if digits else 1
    if style_name.startswith("标题"):
        digits = "".join(c for c in style_name if c.isdigit())
        return int(digits) if digits else 1
    return None


def parse_docx(path: str) -> ParsedDoc:
    """解析 docx，返回结构化段落列表。"""
    doc = DocxDocument(path)
    paragraphs: list[ParagraphMeta] = []

    def collect_paragraphs(container, base_index_offset: int) -> int:
        nonlocal paragraphs
        for p in container.paragraphs:
            text = p.text.strip()
            idx = len(paragraphs)
            style_name = p.style.name if p.style else None
            font, size_pt, bold = _run_font(p)
            pf = p.paragraph_format
            indent_chars = None
            if pf.first_line_indent is not None:
                # python-docx 无法直接读"字符"单位，按近似换算：1 字符 ≈ 字号(pt)
                indent_pt = pf.first_line_indent.pt if hasattr(pf.first_line_indent, "pt") else None
                char_size = size_pt or 12.0
                if indent_pt:
                    indent_chars = round(indent_pt / char_size, 2)
            line_spacing = None
            line_spacing_pt = None
            if pf.line_spacing is not None:
                if hasattr(pf.line_spacing, "pt"):
                    line_spacing_pt = pf.line_spacing.pt
                else:
                    line_spacing = pf.line_spacing
            level = _heading_level(style_name)
            paragraphs.append(
                ParagraphMeta(
                    index=idx,
                    text=text,
                    font=font,
                    size_pt=size_pt,
                    first_line_indent_chars=indent_chars,
                    line_spacing=line_spacing,
                    line_spacing_pt=line_spacing_pt,
                    space_before_pt=(
                        pf.space_before.pt if pf.space_before is not None else None
                    ),
                    space_after_pt=(
                        pf.space_after.pt if pf.space_after is not None else None
                    ),
                    alignment=_alignment_str(pf.alignment),
                    bold=bold,
                    style_name=style_name,
                    is_heading=level is not None,
                    heading_level=level,
                )
            )
        return 0

    collect_paragraphs(doc, 0)
    # 表格内段落（简化：追加到文档段落之后）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect_paragraphs(cell, 0)

    char_count = sum(len(p.text) for p in paragraphs)
    return ParsedDoc(paragraphs=paragraphs, char_count=char_count)


def parse_to_chunks(path: str, max_chars: int = 1200) -> tuple[ParsedDoc, list[list[int]]]:
    """解析并按段落分块，返回 (解析结果, 每个块包含的段落索引列表)。"""
    parsed = parse_docx(path)
    texts = parsed.body_texts()
    chunks = chunk_paragraphs(texts, max_chars=max_chars)
    return parsed, chunks
