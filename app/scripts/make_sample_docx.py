"""生成带典型问题的示例文档（用于演示/自测全流程）。

生成 data/sample.docx，包含：
- 错别字：必段 / 举形 / 布署
- 专有名词不完整：特色社会主义 / 两个百年
- 格式问题：非宋体字体、字号错误、无首行缩进、行距错误
- 句意冗余：「进行加强管理」「非常非常」
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent.parent / "data" / "sample.docx"


def build_sample(path: Path = OUT) -> Path:
    doc = Document()

    # 标题（黑体、三号、居中）
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("关于开展年度工作布署的通知")
    run.font.name = "黑体"
    run.font.size = Pt(16)
    run.bold = True
    h.paragraph_format.line_spacing = Pt(28)

    # 正文1：错别字 + 专有名词不完整 + 无首行缩进 + 字体错误(Calibri)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("各单位必段高度重视本次工作，按照特色社会主义理论的要求，认真贯彻两个百年的战略目标。")
    r1.font.name = "Calibri"
    r1.font.size = Pt(14)
    p1.paragraph_format.line_spacing = 1.0
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 正文2：错别字 + 句意冗余（正确缩进与字体）
    p2 = doc.add_paragraph()
    r2 = p2.add_run("会议将于下周举形，届时请相关负责人出席。我们要进行加强管理，确保各项工作落实到位，做到非常非常细致。")
    r2.font.name = "宋体"
    r2.font.size = Pt(12)
    p2.paragraph_format.first_line_indent = Pt(24)  # 2 字符
    p2.paragraph_format.line_spacing = Pt(28)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 正文3：正常段落（作为对照）
    p3 = doc.add_paragraph()
    r3 = p3.add_run("请各部门于本周五前提交工作计划，逾期不予受理。")
    r3.font.name = "宋体"
    r3.font.size = Pt(12)
    p3.paragraph_format.first_line_indent = Pt(24)
    p3.paragraph_format.line_spacing = Pt(28)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


if __name__ == "__main__":
    p = build_sample()
    print(f"sample docx generated: {p}")
